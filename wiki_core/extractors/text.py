from __future__ import annotations

import csv
import email
import subprocess
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class ExtractedDocument:
    source: str
    source_type: str
    text: str
    units: list[dict[str, object]]
    warnings: list[str]


def _looks_binary(data: bytes) -> bool:
    """Heuristic: NUL bytes or a high fraction of non-text bytes indicate binary."""
    if b"\x00" in data:
        return True
    if not data:
        return False
    sample = data[:4096]
    text_bytes = bytes(range(0x20, 0x7F)) + b"\n\r\t\f\b"
    nontext = sum(1 for b in sample if b not in text_bytes)
    return (nontext / len(sample)) > 0.30


def _read_text(path: Path) -> tuple[str, list[str]]:
    """Read text, but refuse to return the raw bytes of a binary file."""
    raw = path.read_bytes()
    if _looks_binary(raw):
        return "", [f"unsupported_binary_source:{path.suffix or 'no-ext'}"]
    return raw.decode("utf-8", errors="replace"), []


def _extract_pdf(path: Path) -> tuple[str, list[dict[str, object]], list[str]]:
    warnings: list[str] = []
    try:
        output = subprocess.check_output(
            ["pdftotext", "-layout", str(path), "-"], text=True, stderr=subprocess.PIPE
        )
    except FileNotFoundError:
        # Missing system dependency: do NOT return raw bytes (would pollute the index).
        return "", [], ["pdftotext_unavailable:install_poppler"]
    except subprocess.CalledProcessError as exc:
        return "", [], [f"pdftotext_failed:{exc}"]
    units = [{"unit_id": "pdf-full-text", "kind": "pdf_text", "text": output}]
    if not output.strip():
        warnings.append("empty_text_extraction")
    return output, units, warnings


def _extract_csv(path: Path) -> tuple[str, list[dict[str, object]], list[str]]:
    rows: list[str] = []
    units: list[dict[str, object]] = []
    dialect = "excel-tab" if path.suffix.lower() == ".tsv" else "excel"
    with path.open(newline="", encoding="utf-8", errors="replace") as handle:
        reader = csv.reader(handle, dialect=dialect)
        for index, row in enumerate(reader, start=1):
            line = " | ".join(row)
            rows.append(line)
            units.append({"unit_id": f"row-{index}", "kind": "table_row", "row": index, "text": line})
    return "\n".join(rows), units, []


def _extract_xlsx(path: Path) -> tuple[str, list[dict[str, object]], list[str]]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", [], ["openpyxl_missing:pip_install_openpyxl"]
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:  # corrupted file / unexpected format
        return "", [], [f"xlsx_read_failed:{exc}"]
    lines: list[str] = []
    units: list[dict[str, object]] = []
    for sheet in wb.worksheets:
        for r_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = ["" if c is None else str(c) for c in row]
            if not any(cells):
                continue
            line = " | ".join(cells)
            lines.append(line)
            units.append(
                {"unit_id": f"{sheet.title}-row-{r_index}", "kind": "sheet_row", "sheet": sheet.title, "row": r_index, "text": line}
            )
    wb.close()
    return "\n".join(lines), units, []


def _extract_docx(path: Path) -> tuple[str, list[dict[str, object]], list[str]]:
    try:
        from docx import Document
    except ImportError:
        return "", [], ["python_docx_missing:pip_install_python-docx"]
    try:
        doc = Document(str(path))
    except Exception as exc:
        return "", [], [f"docx_read_failed:{exc}"]
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    units = [
        {"unit_id": f"para-{i}", "kind": "doc_paragraph", "ordinal": i, "text": text}
        for i, text in enumerate(paragraphs, start=1)
    ]
    return "\n".join(paragraphs), units, []


def _extract_email(path: Path) -> tuple[str, list[dict[str, object]], list[str]]:
    raw = path.read_bytes()
    try:
        msg = email.message_from_bytes(raw)
    except Exception as exc:
        return "", [], [f"email_parse_failed:{exc}"]
    header_lines = [
        f"{key}: {msg.get(key)}" for key in ("From", "To", "Cc", "Date", "Subject") if msg.get(key)
    ]
    body_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain" and not part.get_filename():
                payload = part.get_payload(decode=True)
                if payload:
                    body_parts.append(payload.decode(part.get_content_charset() or "utf-8", errors="replace"))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            body_parts.append(payload.decode(msg.get_content_charset() or "utf-8", errors="replace"))
    text = "\n".join(header_lines + ["", *body_parts])
    units = [{"unit_id": "email-headers", "kind": "email_headers", "text": "\n".join(header_lines)}]
    if body_parts:
        units.append({"unit_id": "email-body", "kind": "email_body", "text": "\n".join(body_parts)})
    return text, units, []


_SUFFIX_TYPE = {
    ".pdf": "pdf",
    ".csv": "table",
    ".tsv": "table",
    ".xlsx": "spreadsheet",
    ".xlsm": "spreadsheet",
    ".docx": "document",
    ".eml": "email",
    ".mbox": "email",
}


def extract_source(source: str, source_type: str) -> ExtractedDocument:
    path = Path(source).expanduser()
    if not path.exists():
        return ExtractedDocument(source, source_type, "", [], ["source_not_found"])

    # Reconcile a generic source_type (document/email/file) with the real
    # extension, so we do not fall into raw-text reads of binaries (xlsx/docx).
    suffix_type = _SUFFIX_TYPE.get(path.suffix.lower())
    if source_type in {"file", "document", "email", "spreadsheet", "table"} and suffix_type:
        source_type = suffix_type

    if source_type == "pdf":
        text, units, warnings = _extract_pdf(path)
    elif source_type == "spreadsheet":
        text, units, warnings = _extract_xlsx(path)
    elif source_type == "document":
        text, units, warnings = _extract_docx(path)
    elif source_type == "email":
        text, units, warnings = _extract_email(path)
    elif source_type == "table":
        text, units, warnings = _extract_csv(path)
    elif source_type in {"markdown", "file", "text"}:
        text, warnings = _read_text(path)
        units = [{"unit_id": "full-text", "kind": "text", "text": text}]
    else:
        text, warnings = _read_text(path)
        units = [{"unit_id": "full-text", "kind": "text", "text": text}]

    return ExtractedDocument(source, source_type, text, units, warnings)
